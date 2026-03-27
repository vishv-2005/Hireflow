# resume_parser.py - handles extracting text from PDF, DOCX, DOC, and image resumes
# we use PyMuPDF (fitz) for PDFs, python-docx for Word files,
# mammoth for legacy .doc files, and easyocr for image OCR (no external binary needed)

import fitz  # PyMuPDF
import docx
import mammoth
import easyocr
from PIL import Image
import numpy as np
import os
import io

# initialize easyocr reader once (lazy-loaded on first call to avoid startup delay)
_ocr_reader = None

def _get_ocr_reader():
    """Lazy-initialize the easyocr reader so the model is only loaded when needed."""
    global _ocr_reader
    if _ocr_reader is None:
        print("initializing easyocr reader (first time may download models ~100MB)...")
        _ocr_reader = easyocr.Reader(['en'], gpu=False)
        print("easyocr reader ready!")
    return _ocr_reader


# supported file extensions grouped by type
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
DOC_EXTENSIONS = {".doc"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp"}

# all supported extensions combined (useful for filtering in app.py)
ALL_SUPPORTED_EXTENSIONS = PDF_EXTENSIONS | DOCX_EXTENSIONS | DOC_EXTENSIONS | IMAGE_EXTENSIONS

# minimum characters from a PDF page to consider it "text-based"
# if below this threshold, we treat the page as scanned and run OCR
MIN_TEXT_LENGTH_PER_PAGE = 30


def ocr_image(image):
    """
    Run OCR on a PIL Image using easyocr.
    Converts the image to a numpy array for easyocr processing.
    Returns the extracted text as a single string.
    """
    reader = _get_ocr_reader()
    try:
        # convert PIL image to RGB numpy array (easyocr expects this)
        if image.mode not in ("L", "RGB"):
            image = image.convert("RGB")
        img_array = np.array(image)
        
        # run OCR - easyocr returns list of (bbox, text, confidence) tuples
        results = reader.readtext(img_array, detail=0, paragraph=True)
        return "\n".join(results)
    except Exception as e:
        print(f"easyocr error: {e}")
        return ""


def parse_pdf(filepath):
    """
    Opens a PDF and extracts text from every page.
    For scanned/image-based PDFs where text extraction yields very little,
    we fall back to OCR on each page (renders page as image, then runs easyocr).
    """
    all_text = []
    try:
        doc = fitz.open(filepath)
        for page_num, page in enumerate(doc):
            # first try normal text extraction
            page_text = page.get_text().strip()
            
            if len(page_text) >= MIN_TEXT_LENGTH_PER_PAGE:
                # the page has enough text, use it directly
                all_text.append(page_text)
            else:
                # page has very little text — likely a scanned/image PDF
                # render the page as an image and OCR it
                print(f"  page {page_num + 1}: low text ({len(page_text)} chars), running OCR fallback...")
                try:
                    # render at 2x zoom for better OCR accuracy
                    mat = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    ocr_text = ocr_image(img)
                    if ocr_text.strip():
                        all_text.append(ocr_text)
                    elif page_text:
                        # OCR got nothing, but we had some text, keep it
                        all_text.append(page_text)
                except Exception as ocr_err:
                    print(f"  OCR fallback failed on page {page_num + 1}: {ocr_err}")
                    if page_text:
                        all_text.append(page_text)
        doc.close()
    except Exception as e:
        print(f"error reading PDF {filepath}: {e}")
    
    return "\n".join(all_text)


def parse_docx(filepath):
    """
    Opens a DOCX file and grabs text from all paragraphs.
    Also extracts text from tables (which python-docx treats separately).
    """
    text_parts = []
    try:
        doc = docx.Document(filepath)
        # extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # also extract text from tables (resumes often have skills in tables)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
    except Exception as e:
        print(f"error reading DOCX {filepath}: {e}")
    
    return "\n".join(text_parts)


def parse_doc(filepath):
    """
    Opens a legacy .doc file using mammoth.
    Mammoth extracts the raw text content from older Word formats.
    """
    text = ""
    try:
        with open(filepath, "rb") as f:
            result = mammoth.extract_raw_text(f)
            text = result.value
            if result.messages:
                for msg in result.messages:
                    print(f"mammoth warning for {filepath}: {msg}")
    except Exception as e:
        print(f"error reading DOC {filepath}: {e}")
        text = ""
    return text


def parse_image(filepath):
    """
    Uses easyocr to extract text from image files (PNG, JPG, TIFF, BMP, etc.).
    This is essential for scanned resume documents or screenshot resumes.
    easyocr is pure Python — no external Tesseract binary needed!
    """
    text = ""
    try:
        img = Image.open(filepath)
        text = ocr_image(img)
    except Exception as e:
        print(f"error OCR-ing image {filepath}: {e}")
        text = ""
    return text


def parse_resume(filepath):
    """
    Figures out what type of file we're dealing with and calls
    the right parser. Returns a dict with the filename and the raw extracted text.
    """
    filename = os.path.basename(filepath)
    extension = os.path.splitext(filepath)[1].lower()

    if extension in PDF_EXTENSIONS:
        raw_text = parse_pdf(filepath)
    elif extension in DOCX_EXTENSIONS:
        raw_text = parse_docx(filepath)
    elif extension in DOC_EXTENSIONS:
        raw_text = parse_doc(filepath)
    elif extension in IMAGE_EXTENSIONS:
        raw_text = parse_image(filepath)
    else:
        print(f"skipping unsupported file type: {filename}")
        return None

    if not raw_text or not raw_text.strip():
        print(f"warning: no text extracted from {filename}")

    return {
        "filename": filename,
        "raw_text": raw_text or ""
    }
